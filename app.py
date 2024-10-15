from flask import Flask, render_template, request, redirect, url_for, send_file
import pytesseract
from PIL import Image
import os
import docx

app = Flask(__name__)

UPLOAD_FOLDER = 'static/uploads/'
app.config['UPLOAD_FOLDER'] = UPLOAD_FOLDER

# Tesseract executable path
pytesseract.pytesseract.tesseract_cmd = r'C:\Program Files\Tesseract-OCR\tesseract.exe'

# Ensure the upload folder exists
if not os.path.exists(app.config['UPLOAD_FOLDER']):
    os.makedirs(app.config['UPLOAD_FOLDER'])


def extract_text_from_image(image_path):
    # Görüntüyü aç
    img = Image.open(image_path)

    # OCR işlemi ile metni çıkar
    text = pytesseract.image_to_string(img)

    # Metni satırlara ayır (her satırı bir liste öğesi yap)
    lines = text.split('\n')

    # Boş satırları temizle
    cleaned_lines = [line.strip() for line in lines if line.strip()]

    return cleaned_lines


def create_docx_file(text_lines):
    # Yeni bir DOCX dosyası oluştur
    doc = docx.Document()

    # Başlık ekle (opsiyonel)
    doc.add_heading('Extracted OCR Text', 0)

    # Her bir satırı DOCX dosyasına paragraf olarak ekle
    for line in text_lines:
        doc.add_paragraph(line)

    # DOCX dosyasını kaydet
    docx_path = os.path.join(app.config['UPLOAD_FOLDER'], 'extracted_text.docx')
    doc.save(docx_path)

    return docx_path


@app.route('/')
def home():
    return render_template('index.html')


@app.route('/upload', methods=['POST'])
def upload_image():
    if 'image' not in request.files:
        return redirect(request.url)
    file = request.files['image']
    if file.filename == '':
        return redirect(request.url)

    if file:
        # Yüklenen görüntüyü kaydet
        image_path = os.path.join(app.config['UPLOAD_FOLDER'], file.filename)
        file.save(image_path)

        # OCR ile metni çıkar
        extracted_lines = extract_text_from_image(image_path)

        # DOCX dosyasını oluştur
        docx_file_path = create_docx_file(extracted_lines)

        # Sadece dosya adını almak için os.path.basename kullan
        docx_file_name = os.path.basename(docx_file_path)

        return render_template('result.html', extracted_lines=extracted_lines, docx_file_name=docx_file_name,
                               image_path=image_path)


@app.route('/download/<filename>')
def download_file(filename):
    # DOCX dosyasını indirme linki sun
    return send_file(os.path.join(app.config['UPLOAD_FOLDER'], filename), as_attachment=True)


if __name__ == "__main__":
    app.run(debug=True)
